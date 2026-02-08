from docx import Document
from docx.document import Document as DocxDocument
from app.text_extraction.base_extracter import BaseExtracter

class DocExtracter(BaseExtracter):
    MAX_CHARS = 800   

    def extract_text(self) -> list[str]:
        """
        Extract text from the DOCX file, chunking by paragraphs and tables."""
        doc_stream = self.get_file_bytes_stream()
        doc_reader = Document(doc_stream)
        
        para_chunks = self.__create_chunks_by_char_limit(doc_reader)
        table_chunks = self.__get_tables_text_list(doc_reader)
        chunks= para_chunks + table_chunks
        
        if not chunks:
            raise ValueError("The DOCX file contains no extractable content.")
        
        return chunks

    def __get_paragraphs_text_list(self, doc_reader: DocxDocument) -> list[str]:
        chunks = []
        for paragraph in doc_reader.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)
        return chunks
    
    def __get_tables_text_list(self, doc_reader: DocxDocument) -> list[str]:
        tables = []
        for table in doc_reader.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            tables.append("\n".join(rows))
        return tables
    
    def __create_chunks_by_char_limit(self, doc_reader: DocxDocument) -> list[str]:
        text_list=[]
        for para in doc_reader.paragraphs:
            if para.text.strip():
                text_list.append(para.text.strip())
        
        chunks = []
        current = []
        current_len = 0
        for text in text_list:
            text_len = len(text)

            if current_len + text_len > self.MAX_CHARS:
                chunks.append(" ".join(current))
                current = []
                current_len = 0
            current.append(text)
            current_len += text_len

        if current:
            chunks.append(" ".join(current))

        return chunks
